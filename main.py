import sqlite3
import sys

import docx
import requests
from datetime import datetime
from bs4 import BeautifulSoup
from PyQt5.QtWidgets import QApplication, QPushButton, QHBoxLayout, QSizePolicy, QVBoxLayout, QLineEdit, QComboBox
from PyQt5.QtWidgets import QWidget, QMainWindow, QTableWidgetItem, QAbstractItemView
from PyQt5 import QtCore, QtGui, QtWidgets
from docx import Document
from docx.shared import Inches


class Main(QWidget):
    def __init__(self):
        super(Main, self).__init__()
        self.initUI()

    def initUI(self):
        self.setGeometry(300, 300, 400, 400)
        self.setWindowTitle('Конвертор валют')
        self.first_page_layout = QHBoxLayout(self)

        self.btn_stat = QPushButton('Статистика по валютам', self)
        self.btn_stat.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.first_page_layout.addWidget(self.btn_stat)
        self.btn_stat.clicked.connect(self.onClicked_stat)

        self.btn_conver = QPushButton('Конвертор валют', self)
        self.btn_conver.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.first_page_layout.addWidget(self.btn_conver)
        self.btn_conver.clicked.connect(self.onClicked_convert)

    def onClicked_stat(self):
        self.stat_class = Static()
        self.stat_class.show()

    def onClicked_convert(self):
        self.convert_class = Auth()
        self.convert_class.show()


class UI_Static(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(270, 20, 261, 41))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.label.setFont(font)
        self.label.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.label.setTextFormat(QtCore.Qt.AutoText)
        self.label.setAlignment(QtCore.Qt.AlignCenter)
        self.label.setObjectName("label")
        self.tableWidget = QtWidgets.QTableWidget(self.centralwidget)
        self.tableWidget.setGeometry(QtCore.QRect(20, 70, 761, 481))
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.setColumnCount(0)
        self.tableWidget.setRowCount(0)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 26))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label.setText(_translate("MainWindow", "Все валюты"))


class Static(QMainWindow, UI_Static):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.currency = {}
        self.parse()
        self.table_values()
        self.setWindowTitle('Статистика по валютам')

    def parse(self):
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                                 'AppleWebKit/537.36 (KHTML, like Gecko) '
                                 'Chrome/93.0.4577.82 Safari/537.36'}
        url = 'https://finance.rambler.ru/currencies/'

        page = requests.get(url, headers=headers)
        soup = BeautifulSoup(page.content, 'lxml')
        found = soup.find_all('a', class_='finance-currency-table__tr')
        self.currency['Рубль'] = ['RUS', '1', '1', '0', '0']
        class_list = ['finance-currency-table__cell--code', 'finance-currency-table__cell--denomination',
                      'finance-currency-table__cell--value', 'finance-currency-table__cell--change',
                      'finance-currency-table__cell--percent']
        for item in found:
            name = item.find('div', class_='finance-currency-table__cell--currency').text.strip()
            self.currency[name] = []
            for class_ in class_list:
                self.currency[name].append(item.find('div', class_=class_).text.strip())

    def table_values(self):
        self.tableWidget.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.tableWidget.setColumnCount(6)
        self.tableWidget.setHorizontalHeaderLabels(["Валюта", "Код", "Номинал", 'Курс ЦБ', 'Изменения', '%'])
        self.tableWidget.setRowCount(len(self.currency))

        for z, i in enumerate(self.currency):
            self.tableWidget.setItem(z, 0, QTableWidgetItem(i))
            for j in range(len(self.currency[i])):
                self.tableWidget.setItem(z, j + 1, QTableWidgetItem((self.currency[i][j])))


class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(906, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.groupBox = QtWidgets.QGroupBox(self.centralwidget)
        self.groupBox.setGeometry(QtCore.QRect(30, 20, 361, 511))
        self.groupBox.setTitle("")
        self.groupBox.setObjectName("groupBox")
        self.login_label = QtWidgets.QLabel(self.groupBox)
        self.login_label.setGeometry(QtCore.QRect(120, 10, 160, 30))
        font = QtGui.QFont()
        font.setPointSize(16)
        self.login_label.setFont(font)
        self.login_label.setObjectName("login_label")
        self.line_login = QtWidgets.QLineEdit(self.groupBox)
        self.line_login.setGeometry(QtCore.QRect(50, 90, 250, 30))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.line_login.setFont(font)
        self.line_login.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.line_login.setInputMask("")
        self.line_login.setText("")
        self.line_login.setAlignment(QtCore.Qt.AlignCenter)
        self.line_login.setObjectName("line_login")
        self.line_login_p = QtWidgets.QLineEdit(self.groupBox)
        self.line_login_p.setGeometry(QtCore.QRect(50, 160, 250, 30))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.line_login_p.setFont(font)
        self.line_login_p.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.line_login_p.setText("")
        self.line_login_p.setEchoMode(QtWidgets.QLineEdit.Password)
        self.line_login_p.setAlignment(QtCore.Qt.AlignCenter)
        self.line_login_p.setObjectName("line_login_p")
        self.login_btn = QtWidgets.QPushButton(self.groupBox)
        self.login_btn.setGeometry(QtCore.QRect(100, 220, 160, 30))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.login_btn.setFont(font)
        self.login_btn.setObjectName("login_btn")
        self.login_error = QtWidgets.QLabel(self.groupBox)
        self.login_error.setGeometry(QtCore.QRect(30, 330, 271, 91))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.login_error.setFont(font)
        self.login_error.setText("")
        self.login_error.setObjectName("login_error")
        self.groupBox_2 = QtWidgets.QGroupBox(self.centralwidget)
        self.groupBox_2.setGeometry(QtCore.QRect(450, 20, 401, 511))
        self.groupBox_2.setTitle("")
        self.groupBox_2.setObjectName("groupBox_2")
        self.register_label = QtWidgets.QLabel(self.groupBox_2)
        self.register_label.setGeometry(QtCore.QRect(150, 10, 160, 30))
        font = QtGui.QFont()
        font.setPointSize(16)
        self.register_label.setFont(font)
        self.register_label.setObjectName("register_label")
        self.line_register_login = QtWidgets.QLineEdit(self.groupBox_2)
        self.line_register_login.setGeometry(QtCore.QRect(80, 90, 250, 30))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.line_register_login.setFont(font)
        self.line_register_login.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.line_register_login.setInputMask("")
        self.line_register_login.setText("")
        self.line_register_login.setEchoMode(QtWidgets.QLineEdit.Normal)
        self.line_register_login.setAlignment(QtCore.Qt.AlignCenter)
        self.line_register_login.setObjectName("line_register_login")
        self.line_register_p = QtWidgets.QLineEdit(self.groupBox_2)
        self.line_register_p.setGeometry(QtCore.QRect(80, 150, 250, 30))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.line_register_p.setFont(font)
        self.line_register_p.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.line_register_p.setText("")
        self.line_register_p.setEchoMode(QtWidgets.QLineEdit.Password)
        self.line_register_p.setAlignment(QtCore.Qt.AlignCenter)
        self.line_register_p.setObjectName("line_register_p")
        self.line_register_p_2 = QtWidgets.QLineEdit(self.groupBox_2)
        self.line_register_p_2.setGeometry(QtCore.QRect(80, 210, 250, 30))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.line_register_p_2.setFont(font)
        self.line_register_p_2.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.line_register_p_2.setText("")
        self.line_register_p_2.setEchoMode(QtWidgets.QLineEdit.Password)
        self.line_register_p_2.setAlignment(QtCore.Qt.AlignCenter)
        self.line_register_p_2.setObjectName("line_register_p_2")
        self.register_btn = QtWidgets.QPushButton(self.groupBox_2)
        self.register_btn.setGeometry(QtCore.QRect(130, 270, 160, 30))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.register_btn.setFont(font)
        self.register_btn.setObjectName("register_btn")
        self.register_error = QtWidgets.QLabel(self.groupBox_2)
        self.register_error.setGeometry(QtCore.QRect(80, 340, 260, 90))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.register_error.setFont(font)
        self.register_error.setText("")
        self.register_error.setObjectName("register_error")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 906, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.login_label.setText(_translate("MainWindow", "Авторизация"))
        self.line_login.setPlaceholderText(_translate("MainWindow", "Логин"))
        self.line_login_p.setPlaceholderText(_translate("MainWindow", "Пароль"))
        self.login_btn.setText(_translate("MainWindow", "Авторизоваться"))
        self.register_label.setText(_translate("MainWindow", "Регистрация"))
        self.line_register_login.setPlaceholderText(_translate("MainWindow", "Логин"))
        self.line_register_p.setPlaceholderText(_translate("MainWindow", "Пароль"))
        self.line_register_p_2.setPlaceholderText(_translate("MainWindow", "Введите пароль еще раз"))
        self.register_btn.setText(_translate("MainWindow", "Регистрация"))


class Auth(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.con = sqlite3.connect("profiles.db")
        self.cur = self.con.cursor()
        self.bad_simb = set("~!@#£€$¢¥§%^&*/()\\-_+={}[]:;\"'<>,.?")
        self.login_btn.clicked.connect(self.login)
        self.register_btn.clicked.connect(self.register)
        self.setWindowTitle('Окно авторизации')

    def register(self):
        login = self.line_register_login.text()
        p = self.line_register_p.text()
        p2 = self.line_register_p_2.text()
        if login == '' or p == '' or p2 == '':
            self.register_error.setText('Заполнены не все поля \n для ввода')

        elif any((c in self.bad_simb) for c in login):
            self.register_error.setText('Содержаться некорректные \n символы в логине')

        elif p != p2 or any((c in self.bad_simb) for c in p):
            self.register_error.setText('Пароли не совпадают или \n содержат некорректные \n символы')

        elif self.cur.execute("""SELECT id FROM profiles WHERE login = ?""", (login,)).fetchone():
            self.register_error.setText('Такой логин уже занят')

        else:
            try:
                self.cur.execute("""INSERT INTO profiles(login, password) VALUES(?, ?)""", (login, p))
                self.con.commit()
                self.register_error.setText('Успешная регистрация \n Авторизуйтесь')
            except:
                self.con.rollback()
                self.register_error.setText('Ошибка регистрации \n Попробуйте позже')

    def login(self):
        login = self.line_login.text()
        p = self.line_login_p.text()

        if login == '' or p == '':
            self.login_error.setText('Заполнены не все полня для ввода')

        else:
            fetch = self.cur.execute("""SELECT * FROM profiles WHERE login = ?""", (login,)).fetchone()
            if fetch is not None and fetch[2] == p:
                self.login_error.setText('Вы авторизованы')
                user_id = fetch[0]
                self.conv_class = Convert(user_id)
                self.conv_class.show()
                self.con.close()
                self.close()

            elif fetch is None:
                self.login_error.setText('Такого аккаунта не \n существует')
            else:
                self.login_error.setText('Неверный пароль')


class Convert(QWidget):
    def __init__(self, user_id):
        super().__init__()
        self.user_id = user_id
        self.currency = {}
        self.trace_currency()
        self.initUI()
        self.con = sqlite3.connect("profiles.db")
        self.cur = self.con.cursor()

    def initUI(self):
        self.setWindowTitle("Конвертер валют")

        self.main_layout = QHBoxLayout(self)
        self.input_layout = QVBoxLayout(self)
        self.output_layout = QVBoxLayout(self)
        self.currency_layout = QVBoxLayout(self)

        self.input_value = QLineEdit(self)
        self.input_value.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self.input_type = QComboBox(self)
        self.input_type.addItems(self.currency.keys())
        self.input_type.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self.convert_button = QPushButton(self)
        self.convert_button.setText("->")
        self.convert_button.clicked.connect(self.convert)
        self.convert_button.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self.output_value = QLineEdit(self)
        self.output_value.setEnabled(False)
        self.output_value.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self.output_type = QComboBox(self)
        self.output_type.addItems(self.currency.keys())
        self.output_type.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self.transactions_button = QPushButton(self)
        self.transactions_button.setText('Просмотреть историю конвертаций')
        self.transactions_button.clicked.connect(self.onClicked_history)

        self.save_txt_button = QPushButton(self)
        self.save_txt_button.setText('Сохранить историю в файл')
        self.save_txt_button.clicked.connect(self.onClicked_save_file)

        self.input_layout.addWidget(self.input_value)
        self.input_layout.addWidget(self.input_type)
        self.input_layout.addWidget(self.transactions_button)

        self.output_layout.addWidget(self.output_value)
        self.output_layout.addWidget(self.output_type)
        self.output_layout.addWidget(self.save_txt_button)

        self.main_layout.addLayout(self.currency_layout)
        self.main_layout.addLayout(self.input_layout)
        self.main_layout.addWidget(self.convert_button)
        self.main_layout.addLayout(self.output_layout)
        self.setLayout(self.main_layout)

    def convert(self):
        if self.input_value.text() == '' or self.input_value.text().isdigit() is False:
            pass
        else:
            input_ = float(self.input_value.text()) * (float(self.currency[self.input_type.currentText()][2]) / float(
                self.currency[self.input_type.currentText()][1]))
            output = input_ / (float(self.currency[self.output_type.currentText()][2])
                               / float(self.currency[self.output_type.currentText()][1]))
            self.output_value.setText(f"{output:.2f}")

            self.save_history()

    def trace_currency(self):
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) '
                                 'AppleWebKit/537.36 (KHTML, like Gecko) '
                                 'Chrome/93.0.4577.82 Safari/537.36'}
        url = 'https://finance.rambler.ru/currencies/'

        page = requests.get(url, headers=headers)
        soup = BeautifulSoup(page.content, 'lxml')
        found = soup.find_all('a', class_='finance-currency-table__tr')
        self.currency['Рубль'] = ['RUS', '1', '1', '0', '0']
        class_list = ['finance-currency-table__cell--code', 'finance-currency-table__cell--denomination',
                      'finance-currency-table__cell--value', 'finance-currency-table__cell--change',
                      'finance-currency-table__cell--percent']
        for item in found:
            name = item.find('div', class_='finance-currency-table__cell--currency').text.strip()
            self.currency[name] = []
            for class_ in class_list:
                self.currency[name].append(item.find('div', class_=class_).text.strip())

    def onClicked_history(self):
        self.stat_class = History_of_converts(self.user_id)
        self.stat_class.show()

    def onClicked_save_file(self):
        document = docx.Document()
        document.add_heading('История конвертаций', 0)
        self.conn = sqlite3.connect("profiles.db")
        self.cur = self.conn.cursor()
        result = self.cur.execute("""SELECT id_convert, id_first_val, first_sum, id_second_val, second_sum, course, date
                FROM convert WHERE id_user = ?""", (self.user_id,)).fetchall()
        for i in result:
            lst = list(i)
            lst.remove(lst[0])
            first_val = lst[0]
            lst.remove(lst[0])
            second_val = lst[1]
            lst.remove(lst[1])
            result_first_value = self.cur.execute("""SELECT name FROM currency WHERE id = ?""",
                                                  (first_val,)).fetchall()
            result_second_value = self.cur.execute("""SELECT name FROM currency WHERE id = ?""",
                                                   (second_val,)).fetchall()
            tmp = ''.join(list(*result_first_value))
            tmp2 = ''.join(list(*result_second_value))
            lst.insert(0, tmp)
            lst.insert(2, tmp2)
            output = f'{lst[1]} {lst[0]} -> {lst[3]} {lst[2]}. По курсу: {lst[4]}'
            document.add_paragraph(output)
        document.save('history-convert.docx')

    def save_history(self):
        id_first_val = self.cur.execute("""SELECT id FROM currency WHERE name = ?""",
                                        (self.input_type.currentText(),)).fetchone()[0]
        id_second_val = self.cur.execute("""SELECT id FROM currency WHERE name = ?""",
                                         (self.output_type.currentText(),)).fetchone()[0]
        date = datetime.now().date()
        course = f'{round(float(self.currency[self.input_type.currentText()][2]), 2)} к {round(float(self.currency[self.output_type.currentText()][2]), 2)}'

        self.cur.execute("""INSERT INTO convert(id_user, id_first_val, first_sum, id_second_val, second_sum, course,
        date) VALUES(?, ?, ?, ?, ?, ?, ?)""",
                         (self.user_id, id_first_val, self.input_value.text(),
                          id_second_val, self.output_value.text(), course, date))
        self.con.commit()


class Ui_history_of_converts(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.tableWidget = QtWidgets.QTableWidget(self.centralwidget)
        self.tableWidget.setGeometry(QtCore.QRect(30, 20, 731, 521))
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.setColumnCount(0)
        self.tableWidget.setRowCount(0)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))


class History_of_converts(QMainWindow, Ui_history_of_converts):
    def __init__(self, user_id):
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle('История конвертаций')
        self.con = sqlite3.connect("profiles.db")
        self.cur = self.con.cursor()
        self.user_id = user_id
        self.main_func()

    def main_func(self):
        fetch = self.cur.execute("""SELECT id_convert, id_first_val, first_sum, id_second_val, second_sum, course, date
        FROM convert WHERE id_user = ?""", (self.user_id,)).fetchall()
        self.tableWidget.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.tableWidget.setColumnCount(7)
        self.tableWidget.setHorizontalHeaderLabels(["id конвертации",
                                                    "Первая валюта", "Исходная сумма",
                                                    'Вторая валюта', 'Конечная сумма',
                                                    'курс', 'дата'])
        self.tableWidget.setRowCount(len(fetch))
        for i in range(len(fetch)):
            for z in range(len(fetch[i])):
                if z == 1 or z == 3:
                    fetch_cur = self.cur.execute("""SELECT name FROM currency WHERE id = ?""",
                                                 (fetch[i][z],)).fetchone()[0]
                    self.tableWidget.setItem(i, z, QTableWidgetItem(fetch_cur))
                else:
                    if type(fetch[i][z]) is str:
                        self.tableWidget.setItem(i, z, QTableWidgetItem((fetch[i][z])))
                    else:
                        self.tableWidget.setItem(i, z, QTableWidgetItem(str((fetch[i][z]))))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    wnd = Main()
    wnd.show()
    sys.exit(app.exec())
